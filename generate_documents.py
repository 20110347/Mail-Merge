from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Cm
from email.message import EmailMessage
import smtplib
import ssl
import os
from datetime import datetime

# Fuction para calcular la edad
def cal_age(born):
    # Guarda el dia de hoy
    today = datetime.today()
    # Obtiene la edad a partir de restar el año actual menos el de nacimiento, restando 1 o 0 dependiendo de la comparación de los meses y dias
    age = today.year - born.year - \
        ((today.month, today.day) < (born.month, born.day))
    return age

########################################## TXT #########################################################
# Por cada destinatario generara un documento de texto
def generate_txt(recipient, fullTemplate):
    born = datetime.strptime(recipient["fnac"], '%Y-%m-%d')
    edadC = cal_age(born)
    
    # Generar la carpeta si no existe
    results = os.path.join("/txt")
    if not os.path.exists(results):
        os.makedirs(results) 

    # Nombre del Archivo Generado
    file_string = recipient['name'][0:2] + recipient['apellido1'][0:2] + recipient['apellido2'][0:2]
    print(file_string)

    with open("txt/"+file_string+".txt", "w", encoding="utf-8") as file_generated:
        file_generated.write(fullTemplate.format(
            nombre=recipient['name'],
            apellido1=recipient['apellido1'],
            apellido2=recipient['apellido2'],
            cargo=recipient['cargo'],
            empresa=recipient['empresa'],
            calle=recipient['calle'],
            noExt=recipient['noExt'],
            noInt=recipient['noInt'],
            colonia=recipient['colonia'],
            mun=recipient['municipio'],
            estado=recipient['estado'],
            cp=recipient['cp'],
            tel=recipient['tel'],
            fn=recipient['fnac'],
            email=recipient['email'],
            edad=edadC
        ))
    print("\n TXT Generados \n")

########################################## PDF #########################################################

def generate_pdf(recipient, fullTemplate):
    # Por cada destinatario generara un documento de texto
    date = datetime.now()
    date_generated = date.strftime('%Y/%m/%d %H:%M:%S')
    born = datetime.strptime(recipient["fnac"], '%Y-%m-%d')
    edadC = cal_age(born)

    # Generar la carpeta si no existe
    results = os.path.join("/pdf")
    if not os.path.exists(results):
        os.makedirs(results) 

    # Nombre del Archivo Generado
    file_string = recipient['name'][0:2] + recipient['apellido1'][0:2] + recipient['apellido2'][0:2]
    print(file_string)

    # Altura y Anchura tamaño carta
    w, h = letter

    # doc = SimpleDocTemplate("pdf/"+ file_string +".pdf", pagesize=letter,
    #     rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    doc = canvas.Canvas("pdf/"+ file_string +".pdf", pagesize=letter)

    img = ImageReader("GUI/ceti_logo.png")
    img_w, img_h = img.getSize()

    doc.drawImage(img, 50, h - img_h - 25)
    doc.drawString(w - 200, h - 122, date_generated)
    doc.line(50, h - 125 , w - 50, h - 125)

    text = doc.beginText(50, h - 150)
    text.setFont("Times-Roman", 12)

    text.textLines(fullTemplate.format(
        nombre=recipient['name'],
        apellido1=recipient['apellido1'],
        apellido2=recipient['apellido2'],
        cargo=recipient['cargo'],
        empresa=recipient['empresa'],
        calle=recipient['calle'],
        noExt=recipient['noExt'],
        noInt=recipient['noInt'],
        colonia=recipient['colonia'],
        mun=recipient['municipio'],
        estado=recipient['estado'],
        cp=recipient['cp'],
        tel=recipient['tel'],
        fn=recipient['fnac'],
        email=recipient['email'],
        edad=edadC,))
    doc.drawText(text)
    doc.showPage()
    doc.save()
    print("\n PDF Generados \n")

########################################## WORD #########################################################

def generate_docx(recipient, fullTemplate):
    # Por cada destinatario generara un documento de texto
    date = datetime.now()
    date_generated = date.strftime('%Y/%m/%d %H:%M:%S')
    born = datetime.strptime(recipient["fnac"], '%Y-%m-%d')
    edadC = cal_age(born)

    # Generar la carpeta si no existe
    results = os.path.join("/docx")
    if not os.path.exists(results):
        os.makedirs(results)

    # Nombre del Archivo Generado
    file_string = recipient['name'][0:2] + recipient['apellido1'][0:2] + recipient['apellido2'][0:2]
    print(file_string)

    doc = Document()

    doc.add_heading('Carta de ' + recipient['name'], level = 1)
    doc.add_picture('GUI/ceti_logo.png', width=Cm(5))

    doc.add_paragraph(fullTemplate.format(
        nombre=recipient['name'],
        apellido1=recipient['apellido1'],
        apellido2=recipient['apellido2'],
        cargo=recipient['cargo'],
        empresa=recipient['empresa'],
        calle=recipient['calle'],
        noExt=recipient['noExt'],
        noInt=recipient['noInt'],
        colonia=recipient['colonia'],
        mun=recipient['municipio'],
        estado=recipient['estado'],
        cp=recipient['cp'],
        tel=recipient['tel'],
        fn=recipient['fnac'],
        email=recipient['email'],
        edad=edadC,))
    doc.add_paragraph("Generado la fecha de: " + date_generated)
    doc.save("docx/"+ file_string +".docx")
    print("\n Docx Generados \n")

########################################## MAILS #########################################################

def generate_mails(recipient, fullTemplate):
    date = datetime.now()
    date_generated = date.strftime('%Y/%m/%d %H:%M:%S')
    born = datetime.strptime(recipient["fnac"], '%Y-%m-%d')
    edadC = cal_age(born)

    # Nombre del Archivo Generado
    file_string = recipient['name'][0:2] + recipient['apellido1'][0:2] + recipient['apellido2'][0:2]
    print(file_string)

    msg = EmailMessage()

    message = (fullTemplate.format(
        nombre=recipient['name'],
        apellido1=recipient['apellido1'],
        apellido2=recipient['apellido2'],
        cargo=recipient['cargo'],
        empresa=recipient['empresa'],
        calle=recipient['calle'],
        noExt=recipient['noExt'],
        noInt=recipient['noInt'],
        colonia=recipient['colonia'],
        mun=recipient['municipio'],
        estado=recipient['estado'],
        cp=recipient['cp'],
        tel=recipient['tel'],
        fn=recipient['fnac'],
        email=recipient['email'],
        edad=edadC,))
        
    password = " PASSWORD OF YOUR GMAIL "
    msg['From'] = " TYPE YOUR EMAIL HERE "
    msg['To'] = recipient['email']
    msg['Subject'] = file_string + "Renuncia, IF YOU SEE THIS MESSAGE, Its a homework for my school"
    msg.set_content(message)

    context = ssl.create_default_context()

    with smtplib.SMTP_SSL('smtp.gmail.com', 465, context=context) as smtp:
        smtp.login("pyrop59@gmail.com", password)
        smtp.sendmail("pyrop59@gmail.com", recipient["email"], msg.as_string())
    print("\n Se enviaron los correos \n")